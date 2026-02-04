import os
from tqdm.auto import tqdm
from doc2docx import convert
from docx import Document
from storage import Storage
import logging

def read_docx(file_path):
    """Read and extract text from a DOCX file."""
    try:
        doc = Document(file_path)
        return '\n'.join(para.text for para in doc.paragraphs)
    except Exception as e:
        logging.error(f"Failed to read DOCX file at {file_path}: {e}")
        return None

def get_documents(folder_path):
    """Retrieve and process documents from a folder, storing them in a database if not already present."""
    #base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    folder_path = os.path.abspath(folder_path)
    # if folder_path is None:
    #     folder_path = os.path.join(base_dir, "3GPP-Release18", "Documents")
    # storage_path = os.path.join(base_dir, "3GPP-Release18", storage_name)
    # storage = Storage(storage_path)
    # storage.create_dataset(dataset_name)
    
    document_ds = []
    #file_list = []
    # Check and convert .doc files to .docx
    #convert_docs_to_docx(folder_path)

    # Prepare list of .docx files for processing
    #file_list = [f for f in os.listdir(folder_path) if valid_file(f, series_list)]
    # Process each document
    # for filename in file_list:
    #     file_path = os.path.join(folder_path, filename)
    #     process_document(file_path, filename, storage, document_ds, dataset_name)

    # storage.close()
    # return document_ds
    for filename in os.listdir(folder_path):
        if not filename.endswith(".docx") or filename.startswith("~$"):
            continue

        file_path = os.path.join(folder_path, filename)

        try:
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            logging.error(f"Failed to read {file_path}: {e}")
            continue

        document_ds.append({
            "id": filename,
            "text": text,
            "source": filename
        })

    return document_ds

def convert_docs_to_docx(folder_path):
    """Convert .doc files in a folder to .docx format if any."""
    has_doc = any(f.endswith('.doc') for f in os.listdir(folder_path))
    if has_doc:
        convert(folder_path)

def valid_file(filename, series_list):
    """Check if a file should be processed based on its name and series list."""
    return filename.endswith(".docx") and not filename.startswith("~$") and (not filename[:2].isnumeric() or int(filename[:2]) in series_list)

def process_document(file_path, filename, storage, document_ds, dataset_name):
    """Process a single document file."""
    if storage.is_id_in_dataset(dataset_name, filename):
        data_dict = storage.get_dict_by_id(dataset_name, filename)
        document_ds.append(data_dict)
    else:
        content = read_docx(file_path)
        if content:
            data_dict = {'id': filename, 'text': content, 'source': filename}
            document_ds.append(data_dict)
            storage.insert_dict(dataset_name, data_dict)

